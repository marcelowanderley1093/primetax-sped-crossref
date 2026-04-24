"""
Gerador de parecer formal em Word (CLAUDE.md §9).

Comando CLI: `primetax-sped parecer <cnpj> <ano> --tese <codigo>`

Saída: documento .docx assinável pelo consultor responsável, consolidando
oportunidades identificadas para uma tese específica (tema-69, insumos,
retencoes, imobilizado, etc).

Estrutura do parecer:
  1. Cabeçalho Primetax (identidade visual §3: cinza #53565A, teal #008C95)
  2. Identificação do cliente e do trabalho (CNPJ, AC, tese, data)
  3. Resumo executivo com tabela de achados
  4. Fundamentação legal (dispositivos da tese)
  5. Análise individualizada de cada oportunidade com evidência rastreável
  6. Conclusão e recomendação
  7. Espaço para assinatura do consultor responsável

O parecer é defensável — cada achado carrega base legal e evidência até a
linha do arquivo SPED de origem (princípio 1 do §4 — rastreabilidade absoluta).
"""

from __future__ import annotations

import json
import os
from datetime import datetime
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from src.db.repo import Repositorio

# Identidade visual Primetax (§3)
_TEAL = RGBColor(0x00, 0x8C, 0x95)
_CINZA = RGBColor(0x53, 0x56, 0x5A)
_BRANCO = RGBColor(0xFF, 0xFF, 0xFF)

# Caminho padrão do logo Primetax (opcional — usado se o arquivo existir).
# Pode ser sobrescrito via variável de ambiente PRIMETAX_LOGO_PATH.
_LOGO_PADRAO = Path("data/identidade/primetax-logo.png")


def _caminho_logo() -> Path | None:
    """Retorna o caminho do logo se existir; None caso contrário."""
    candidato = Path(os.environ.get("PRIMETAX_LOGO_PATH", str(_LOGO_PADRAO)))
    return candidato if candidato.exists() else None


# --------------------------------------------------------------------
# Mapa de teses — cada tese agrupa um conjunto de cruzamentos e a base
# legal que a sustenta (§4 princípio 4 — versionamento com dispositivo).
# --------------------------------------------------------------------

_TESES: dict[str, dict] = {
    "tema-69": {
        "nome": "Tese 69 — Exclusão do ICMS da base do PIS/COFINS",
        "codigos_regras": ["CR-07", "CR-08", "CR-09", "CR-26"],
        "dispositivos_legais": [
            "STF, RE 574.706/PR — Tema 69 de repercussão geral (trânsito 16/09/2021)",
            "Parecer SEI 7698/2021/ME — PGFN (operacionalização da Tese 69)",
            "Lei 10.522/2002, art. 19-A — vinculação da RFB a decisões do STF",
            "Lei 10.637/2002, art. 1º — PIS/PASEP não-cumulativo",
            "Lei 10.833/2003, art. 1º — COFINS não-cumulativo",
            "Lei 5.172/1966 (CTN), art. 168, I — prazo prescricional quinquenal",
        ],
        "descricao_tese": (
            "O Supremo Tribunal Federal, ao julgar o RE 574.706/PR em sede de "
            "repercussão geral (Tema 69), fixou o entendimento de que o ICMS não "
            "compõe a base de cálculo para a incidência do PIS e da COFINS. A "
            "modulação dos efeitos estabeleceu como marco a data de 15/03/2017 "
            "para os contribuintes sem ação judicial transitada em julgado até "
            "essa data. O Parecer SEI 7698/2021/ME da PGFN e o art. 19-A da Lei "
            "10.522/2002 vinculam a RFB a essa orientação. O presente parecer "
            "identifica, na EFD-Contribuições do cliente, registros em que o "
            "ICMS destacado compõe indevidamente a base de PIS/COFINS, com "
            "detalhamento linha a linha e cálculo do crédito a recuperar."
        ),
    },
    "insumos": {
        "nome": "Créditos de Insumos — Conceito de Essencialidade (REsp 1.221.170/PR)",
        "codigos_regras": ["CR-11", "CR-12"],
        "dispositivos_legais": [
            "STJ, REsp 1.221.170/PR — Tema 779 de recursos repetitivos",
            "IN RFB 2.121/2022 — conceito de insumo para PIS/COFINS",
            "Lei 10.637/2002, art. 3º, II — direito a crédito sobre insumos",
            "Lei 10.833/2003, art. 3º, II — direito a crédito sobre insumos",
            "Lei 5.172/1966 (CTN), art. 168, I — prazo prescricional quinquenal",
        ],
        "descricao_tese": (
            "O Superior Tribunal de Justiça, no julgamento do REsp 1.221.170/PR "
            "(Tema 779), firmou o entendimento de que o conceito de insumo para "
            "fins de creditamento de PIS/COFINS deve ser aferido à luz dos "
            "critérios da essencialidade e relevância para a atividade "
            "econômica do contribuinte, afastando as restrições restritivas das "
            "IN SRF 247/2002 e 404/2004. A IN RFB 2.121/2022 acomodou o "
            "entendimento administrativamente. O presente parecer identifica "
            "classificações de itens (TIPO_ITEM) e CSTs que podem configurar "
            "oportunidade de reclassificação para habilitar creditamento."
        ),
    },
    "retencoes": {
        "nome": "Retenções na Fonte — PIS/COFINS/CSLL",
        "codigos_regras": ["CR-16", "CR-17", "CR-18", "CR-19"],
        "dispositivos_legais": [
            "Lei 10.833/2003, art. 30, 33 e 34 — retenções PIS/COFINS/CSLL",
            "Lei 9.430/1996, arts. 64 e 65 — retenções IRPJ/CSLL/PIS/COFINS",
            "IN RFB 1.234/2012 — retenções por órgãos públicos",
            "Lei 5.172/1966 (CTN), art. 168, I — prazo prescricional quinquenal",
        ],
        "descricao_tese": (
            "As retenções na fonte de PIS, COFINS, CSLL e IRRF, quando "
            "efetuadas por órgãos públicos ou por pessoas jurídicas que contratam "
            "serviços (segurança, limpeza, manutenção, etc.), geram direito a "
            "compensação pelo contribuinte retido. Valores retidos declarados nos "
            "registros F600/F700/F800 da EFD-Contribuições e não compensados na "
            "apuração M200/M600, ou ainda registrados nos blocos 1300/1700 com "
            "prazo prescricional expirando, configuram oportunidade de recuperação."
        ),
    },
    "imobilizado": {
        "nome": "Créditos sobre Ativo Imobilizado e Depreciação",
        "codigos_regras": ["CR-14", "CR-15", "CR-22", "CR-23", "CR-27", "CR-35"],
        "dispositivos_legais": [
            "Lei 10.637/2002, art. 3º, VI — créditos sobre depreciação",
            "Lei 10.833/2003, art. 3º, VI — créditos sobre depreciação",
            "Lei 11.774/2008, art. 1º — créditos acelerados sobre imobilizado",
            "IN RFB 2.121/2022 — critérios para creditamento sobre imobilizado",
        ],
        "descricao_tese": (
            "A aquisição de bens para o ativo imobilizado destinados à "
            "produção de bens ou prestação de serviços gera crédito de PIS/COFINS "
            "sobre os encargos de depreciação (art. 3º, VI das Leis 10.637/2002 e "
            "10.833/2003) ou pela alternativa da aquisição (§14 do art. 3º da Lei "
            "10.833/2003). O presente parecer identifica: (i) classificações "
            "indevidas em F120/F130; (ii) bens operacionais que não geraram "
            "crédito; (iii) reconhecimento de CIAP na EFD ICMS/IPI sem aproveitamento "
            "correlato na EFD-Contribuições."
        ),
    },
    "prescricao-quinquenal": {
        "nome": "Prescrição Quinquenal — Retenções e Créditos em Risco de Perda",
        "codigos_regras": ["CR-19"],
        "dispositivos_legais": [
            "Lei 5.172/1966 (CTN), art. 168, I — prazo decadencial de 5 anos",
            "Lei 5.172/1966 (CTN), art. 150, §4º — homologação tácita",
            "STJ, REsp 1.002.932/SP — termo inicial do prazo para compensação",
            "Lei 9.250/1995, art. 39 — incidência de SELIC sobre restituição",
            "Lei 10.833/2003, art. 30, 33 e 34 — retenções PIS/COFINS/CSLL",
        ],
        "descricao_tese": (
            "O Código Tributário Nacional estabelece, em seu art. 168, I, prazo "
            "de cinco anos contados da extinção do crédito tributário para que "
            "o contribuinte pleiteie a restituição ou compensação de valores "
            "pagos indevidamente. Retenções na fonte declaradas nos blocos "
            "1300/1700 da EFD-Contribuições com data de aproveitamento "
            "(PR_REC_RET) superior a 5 anos estão em risco de perda por "
            "prescrição. O presente parecer mapeia esses valores com urgência, "
            "indicando o horizonte temporal para protocolo de PER/DCOMP antes "
            "da consumação do prazo prescricional."
        ),
    },
    "lei-14789-subvencoes": {
        "nome": "Subvenções para Investimento — Regime da Lei 14.789/2023",
        "codigos_regras": ["CR-46"],
        "dispositivos_legais": [
            "Lei 14.789/2023 — novo regime de subvenções para investimento",
            "Lei Complementar 160/2017 — conversão de benefícios ICMS",
            "Lei 12.973/2014, art. 30 — tratamento tributário anterior (revogado)",
            "IN RFB 2.214/2024 — regulamentação da Lei 14.789/2023",
            "ADE Cofis 02/2026 — Manual do Leiaute 12 da ECF",
        ],
        "descricao_tese": (
            "A Lei 14.789/2023 reformulou o regime tributário das subvenções "
            "para investimento, introduzindo mecanismo de habilitação e "
            "concessão de crédito fiscal em substituição à anterior exclusão "
            "da base do IRPJ/CSLL prevista no art. 30 da Lei 12.973/2014 "
            "(revogado). O contribuinte deve declarar os benefícios fiscais "
            "nos registros X480/X485 da ECF e pleitear exclusões "
            "correspondentes no e-Lalur (Bloco M300). O presente parecer "
            "identifica benefícios declarados sem contrapartida de exclusão, "
            "configurando oportunidade de retificação e redução de base de "
            "IRPJ/CSLL dentro dos limites do novo regime."
        ),
    },
    "compensacao-prejuizos": {
        "nome": "Compensação de Prejuízos Fiscais — Parte B do e-Lalur",
        "codigos_regras": ["CR-45"],
        "dispositivos_legais": [
            "Lei 9.065/1995, art. 15 — compensação limitada a 30% do lucro",
            "Lei 12.973/2014, art. 67 — controle da Parte B do e-Lalur",
            "IN RFB 1.700/2017, art. 100 — prazo e forma de aproveitamento",
            "ADE Cofis 02/2026 §12 — registros M500/M510 da ECF",
        ],
        "descricao_tese": (
            "A Parte B do e-Lalur (registros M500 da ECF) controla saldos de "
            "adições temporárias e exclusões futuras que ainda podem impactar "
            "a apuração do IRPJ/CSLL em exercícios subsequentes. Contas com "
            "saldo final não-zero (SD_FIM_LAL > 0) e sem movimento no período "
            "(VL_LCTO_PARTE_A = 0 e VL_LCTO_PARTE_B = 0) em natureza de "
            "exclusão futura (IND_SD_FIM_LAL = 'D') configuram valores "
            "pendentes de aproveitamento — prejuízos fiscais próximos do "
            "limite prescricional, provisões revertidas sem exclusão "
            "correspondente e subvenções esquecidas (pré-Lei 14.789). O "
            "presente parecer identifica contas estagnadas com valor "
            "significativo para acionamento imediato."
        ),
    },
    "lei-do-bem": {
        "nome": "Lei do Bem — Inovação Tecnológica e Desenvolvimento (Lei 11.196/2005)",
        "codigos_regras": ["CR-49"],
        "dispositivos_legais": [
            "Lei 11.196/2005, arts. 17 a 26 — incentivos à inovação tecnológica",
            "Decreto 5.798/2006 — regulamentação dos incentivos da Lei do Bem",
            "IN RFB 1.187/2011 — disciplina dos benefícios fiscais",
            "ADE Cofis 02/2026 §20 — registro X460 da ECF",
        ],
        "descricao_tese": (
            "A Lei 11.196/2005 (Lei do Bem), em seus arts. 17 a 26, concede "
            "incentivos fiscais a pessoas jurídicas tributadas pelo Lucro Real "
            "que realizem pesquisa tecnológica e desenvolvimento de inovação "
            "tecnológica. Os principais incentivos são: (i) exclusão adicional "
            "de 60% a 80% dos dispêndios no cálculo do IRPJ/CSLL (art. 19); "
            "(ii) depreciação/amortização integral dos bens e intangíveis "
            "usados em P&D no próprio ano-calendário (art. 17); (iii) "
            "acréscimo de 20% por pesquisador exclusivamente contratado para "
            "P&D (art. 19-A). O presente parecer identifica dispêndios com "
            "inovação tecnológica declarados no registro X460 da ECF sem "
            "correspondente exclusão no M300 (Parte A do e-Lalur), "
            "configurando oportunidade de retificação e recuperação de IRPJ/CSLL."
        ),
    },
    "creditos-extemporaneos": {
        "nome": "Créditos Extemporâneos — Registros 1101/1501 e Ajustes ECD",
        "codigos_regras": ["CR-25", "CR-41"],
        "dispositivos_legais": [
            "Lei 10.637/2002, art. 3º, §4º — crédito extemporâneo PIS",
            "Lei 10.833/2003, art. 3º, §4º — crédito extemporâneo COFINS",
            "IN RFB 2.121/2022 — aproveitamento de créditos extemporâneos",
            "CFC, ITG 2000 (R1) — ajustes contábeis em exercícios anteriores",
            "IN RFB 2.003/2021 — registros I200 com IND_LCTO='X'",
        ],
        "descricao_tese": (
            "Os arts. 3º, §4º das Leis 10.637/2002 e 10.833/2003 permitem o "
            "aproveitamento de créditos de PIS/COFINS em períodos posteriores "
            "ao de origem, desde que dentro do prazo decadencial e "
            "devidamente escriturados nos registros 1101/1501 (atos "
            "extemporâneos) da EFD-Contribuições. A ITG 2000 (R1) do CFC, "
            "refletida no leiaute 9 da ECD via I200.IND_LCTO='X', formaliza "
            "o suporte contábil desses lançamentos. O presente parecer "
            "identifica: (i) saldos acumulados em 1100/1500 que podem "
            "comportar créditos extemporâneos ainda não aproveitados; (ii) "
            "lançamentos contábeis extemporâneos (I200 com IND_LCTO='X') "
            "sem contrapartida em 1101/1501, sinalizando inconsistência "
            "entre os livros contábil e fiscal."
        ),
    },
}


# --------------------------------------------------------------------
# Helpers de formatação
# --------------------------------------------------------------------

def _moeda(v) -> str:
    try:
        return f"R$ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return str(v)


def _paragrafo_titulo(doc: Document, texto: str, nivel: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.color.rgb = _TEAL if nivel == 1 else _CINZA
    run.font.size = Pt(16 if nivel == 1 else 13)
    p.paragraph_format.space_before = Pt(18 if nivel == 1 else 12)
    p.paragraph_format.space_after = Pt(6)


def _paragrafo_corpo(doc: Document, texto: str, *, negrito: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texto)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = _CINZA
    run.font.bold = negrito


def _paragrafo_citacao(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(texto)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = _CINZA
    run.font.italic = True


def _shade_cell(cell, hex_rgb: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_rgb)
    tc_pr.append(shd)


# --------------------------------------------------------------------
# Função pública
# --------------------------------------------------------------------

def gerar(
    repo: Repositorio,
    ano_calendario: int,
    *,
    tese: str,
    destino: Path,
) -> Path:
    """
    Gera parecer Word para a tese indicada.

    Args:
        repo: Repositório do cliente.
        ano_calendario: Ano-calendário do diagnóstico.
        tese: Código da tese (tema-69, insumos, retencoes, imobilizado).
        destino: Caminho de saída do arquivo .docx.

    Returns:
        Caminho do arquivo gerado.

    Raises:
        ValueError: se a tese não for reconhecida.
    """
    if tese not in _TESES:
        teses_disponiveis = ", ".join(sorted(_TESES))
        raise ValueError(
            f"Tese '{tese}' não reconhecida. Disponíveis: {teses_disponiveis}"
        )

    spec = _TESES[tese]

    conn = repo.conexao()
    try:
        todas_ops = repo.consultar_oportunidades(conn, repo.cnpj, ano_calendario)
        ctx = repo.consultar_sped_contexto(conn)
    finally:
        conn.close()

    # Filtra oportunidades pelos códigos de regra da tese
    ops_tese = [
        op for op in todas_ops
        if op.get("codigo_regra") in spec["codigos_regras"]
    ]

    doc = Document()
    _montar_cabecalho(doc, repo.cnpj, ano_calendario, spec, ctx)
    _montar_descricao_tese(doc, spec)
    _montar_fundamentacao_legal(doc, spec)
    _montar_resumo_executivo(doc, ops_tese)
    _montar_analise_individualizada(doc, ops_tese)
    _montar_conclusao(doc, ops_tese, spec)
    _montar_assinatura(doc)

    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destino))
    return destino


# --------------------------------------------------------------------
# Seções do documento
# --------------------------------------------------------------------

def _montar_cabecalho(
    doc: Document, cnpj: str, ano: int, spec: dict, ctx: dict | None,
) -> None:
    # Logo Primetax (opcional — embedded se disponível em data/identidade/)
    logo = _caminho_logo()
    if logo is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        try:
            run.add_picture(str(logo), height=Cm(1.5))
        except Exception:
            # Se o arquivo existir mas estiver corrompido, caímos no cabeçalho textual.
            pass
        p.paragraph_format.space_after = Pt(6)

    # Faixa superior textual (sempre presente — complementa o logo ou substitui se ausente)
    p = doc.add_paragraph()
    run = p.add_run("PRIMETAX SOLUTIONS")
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _TEAL
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    run = p.add_run("Consultoria Tributária — Recuperação de Créditos Federais")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = _CINZA
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(24)

    # Título do parecer
    p = doc.add_paragraph()
    run = p.add_run("PARECER TÉCNICO DE RECUPERAÇÃO TRIBUTÁRIA")
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _CINZA
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    run = p.add_run(spec["nome"])
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = _TEAL
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)

    # Tabela de identificação
    tabela = doc.add_table(rows=4, cols=2)
    tabela.autofit = False
    tabela.columns[0].width = Cm(5.0)
    tabela.columns[1].width = Cm(11.0)

    cnpj_formatado = f"{cnpj[:2]}.{cnpj[2:5]}.{cnpj[5:8]}/{cnpj[8:12]}-{cnpj[12:]}"
    linhas = [
        ("CNPJ do contribuinte", cnpj_formatado),
        ("Ano-calendário", str(ano)),
        ("Tese analisada", spec["nome"]),
        ("Data de emissão", datetime.now().strftime("%d/%m/%Y")),
    ]
    for idx, (chave, valor) in enumerate(linhas):
        tabela.cell(idx, 0).text = chave
        tabela.cell(idx, 1).text = valor
        tabela.cell(idx, 0).paragraphs[0].runs[0].font.bold = True
        for cell in tabela.rows[idx].cells:
            cell.paragraphs[0].runs[0].font.name = "Calibri"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.color.rgb = _CINZA

    doc.add_paragraph()


def _montar_descricao_tese(doc: Document, spec: dict) -> None:
    _paragrafo_titulo(doc, "1. Fundamentação Jurisprudencial", nivel=1)
    _paragrafo_corpo(doc, spec["descricao_tese"])


def _montar_fundamentacao_legal(doc: Document, spec: dict) -> None:
    _paragrafo_titulo(doc, "2. Dispositivos Legais Aplicáveis", nivel=1)
    for dispositivo in spec["dispositivos_legais"]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(dispositivo)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = _CINZA


def _montar_resumo_executivo(doc: Document, ops: list[dict]) -> None:
    _paragrafo_titulo(doc, "3. Resumo Executivo", nivel=1)

    if not ops:
        _paragrafo_corpo(
            doc,
            "O diagnóstico automatizado realizado sobre os arquivos SPED do "
            "contribuinte não identificou, para a tese analisada, oportunidades "
            "de recuperação que atendessem aos critérios objetivos de disparo "
            "do sistema. Esta ausência de achados automatizados não exclui "
            "eventuais fragilidades cuja identificação demande análise manual "
            "complementar pelo auditor responsável.",
        )
        return

    total_cons = sum(
        Decimal(str(o.get("valor_impacto_conservador") or 0)) for o in ops
    )
    total_max = sum(
        Decimal(str(o.get("valor_impacto_maximo") or 0)) for o in ops
    )

    _paragrafo_corpo(
        doc,
        f"O diagnóstico identificou {len(ops)} oportunidade(s) de recuperação "
        f"tributária relacionada(s) à tese em análise. A tabela abaixo sintetiza "
        f"os achados, ordenados por severidade.",
    )

    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = "Light Grid Accent 1"
    hdr = tabela.rows[0].cells
    hdr[0].text = "Regra"
    hdr[1].text = "Severidade"
    hdr[2].text = "Impacto Conservador"
    hdr[3].text = "Impacto Máximo"
    for c in hdr:
        _shade_cell(c, "008C95")
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = _BRANCO
                r.font.bold = True
                r.font.size = Pt(10)

    ordem_sev = {"alto": 0, "medio": 1, "baixo": 2}
    ops_ord = sorted(
        ops,
        key=lambda o: (
            ordem_sev.get(o.get("severidade", "baixo"), 9),
            o.get("codigo_regra", ""),
        ),
    )
    for op in ops_ord:
        linha = tabela.add_row().cells
        linha[0].text = op.get("codigo_regra", "")
        linha[1].text = (op.get("severidade") or "").upper()
        linha[2].text = _moeda(op.get("valor_impacto_conservador") or 0)
        linha[3].text = _moeda(op.get("valor_impacto_maximo") or 0)
        for c in linha:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)
                    r.font.color.rgb = _CINZA

    doc.add_paragraph()
    _paragrafo_corpo(
        doc,
        f"Impacto total conservador estimado: {_moeda(total_cons)}. "
        f"Impacto total máximo estimado: {_moeda(total_max)}. "
        f"Os valores são estimativas de base apurada pelo sistema; o cálculo "
        f"definitivo do crédito a recuperar, com correção pela taxa SELIC (Lei "
        f"9.250/1995, art. 39) e aplicação das alíquotas cabíveis por período, "
        f"é item do trabalho posterior à homologação deste parecer pelo cliente.",
        negrito=True,
    )


def _montar_analise_individualizada(doc: Document, ops: list[dict]) -> None:
    if not ops:
        return

    _paragrafo_titulo(doc, "4. Análise Individualizada dos Achados", nivel=1)

    for idx, op in enumerate(ops, start=1):
        _paragrafo_titulo(
            doc,
            f"4.{idx}  {op.get('codigo_regra', '')} — "
            f"{(op.get('severidade') or '').upper()}",
            nivel=2,
        )

        _paragrafo_corpo(doc, op.get("descricao", ""))

        # Evidência: arquivo e linha do SPED
        try:
            evs = json.loads(op.get("evidencia_json", "[]"))
        except Exception:
            evs = []
        if evs:
            primeiro = evs[0] if isinstance(evs, list) else evs
            if isinstance(primeiro, dict):
                arquivo = primeiro.get("arquivo") or primeiro.get("arquivo_origem") or "—"
                linha = primeiro.get("linha") or primeiro.get("linha_arquivo") or "—"
                _paragrafo_citacao(
                    doc,
                    f"Evidência: arquivo '{Path(str(arquivo)).name}', linha {linha}.",
                )

        # Impacto individual
        cons = op.get("valor_impacto_conservador") or 0
        maxi = op.get("valor_impacto_maximo") or 0
        if float(cons or 0) > 0 or float(maxi or 0) > 0:
            _paragrafo_corpo(
                doc,
                f"Impacto estimado: conservador {_moeda(cons)} — "
                f"máximo {_moeda(maxi)}.",
                negrito=True,
            )


def _montar_conclusao(doc: Document, ops: list[dict], spec: dict) -> None:
    _paragrafo_titulo(doc, "5. Conclusão e Recomendação", nivel=1)

    if not ops:
        _paragrafo_corpo(
            doc,
            "Diante da ausência de achados automatizados para a tese em "
            "análise, recomenda-se o arquivamento deste relatório como "
            "documentação do procedimento de diagnóstico realizado. Caso "
            "surjam dados complementares ou novas escriturações sejam "
            "transmitidas, o diagnóstico poderá ser refeito.",
        )
        return

    _paragrafo_corpo(
        doc,
        "Diante dos achados consolidados neste parecer, recomenda-se: (i) "
        "homologação interna do diagnóstico por responsável técnico do "
        "contribuinte; (ii) retificação das escriturações pertinentes com os "
        "ajustes identificados; (iii) protocolo de pedido administrativo de "
        "restituição/compensação (PER/DCOMP) pelos valores pagos a maior no "
        "quinquênio (art. 168, I do CTN); (iv) monitoramento dos processos "
        "administrativos e, se necessário, judicialização com apoio do "
        "departamento jurídico da Primetax Solutions. Os valores apresentados "
        "carecem de validação final com base na alíquota e regime efetivamente "
        "aplicados a cada fato gerador no período.",
    )


def _montar_assinatura(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    _paragrafo_corpo(
        doc,
        f"Emitido em {datetime.now().strftime('%d de %B de %Y')}, pela equipe "
        f"técnica da Primetax Solutions.",
    )

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = _CINZA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Consultor Responsável")
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _CINZA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Primetax Solutions — CRC/OAB nº ____________")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = _CINZA
