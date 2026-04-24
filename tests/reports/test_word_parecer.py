"""
Testes do gerador de parecer Word (CLAUDE.md §9).

Cobertura:
  - Tese inválida levanta ValueError.
  - Parecer sem achados gera documento com seção "arquivamento".
  - Parecer com achados gera tabela de resumo + análise individualizada.
  - Integração end-to-end: import → diagnose → parecer.
"""

from pathlib import Path

import pytest
from docx import Document

from src.crossref.engine import Motor
from src.db.repo import Repositorio
from src.parsers import efd_contribuicoes
from src.reports import word_parecer


def _texto_concatenado(doc: Document) -> str:
    """Concatena todos os parágrafos + células de tabelas do documento."""
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                partes.append(cell.text)
    return "\n".join(partes)


class TestTesesInvalida:
    def test_tese_desconhecida_levanta_value_error(self, tmp_path):
        db_dir = tmp_path / "db"
        repo = Repositorio("00000000000100", 2025, base_dir=db_dir)
        repo.criar_banco()

        with pytest.raises(ValueError, match="não reconhecida"):
            word_parecer.gerar(
                repo, 2025, tese="tese-inexistente", destino=tmp_path / "x.docx"
            )


class TestTesesDisponiveis:
    def test_teses_carros_chefe_presentes(self):
        """As teses carros-chefe da Primetax devem estar mapeadas."""
        assert {
            "tema-69", "insumos", "retencoes", "imobilizado",
            "prescricao-quinquenal", "lei-14789-subvencoes",
            "compensacao-prejuizos", "creditos-extemporaneos",
            "lei-do-bem",
        }.issubset(set(word_parecer._TESES))

    def test_cada_tese_tem_dispositivos_e_regras(self):
        """Princípio 4 (§4): toda regra carrega base legal."""
        for codigo, spec in word_parecer._TESES.items():
            assert spec["dispositivos_legais"], f"{codigo}: sem dispositivos legais"
            assert spec["codigos_regras"], f"{codigo}: sem cruzamentos mapeados"
            assert spec["nome"], f"{codigo}: sem nome"
            assert spec["descricao_tese"], f"{codigo}: sem descrição"

    def test_tese_prescricao_gera_parecer_valido(self, tmp_path):
        """Smoke test de geração para uma das teses novas (prescrição-quinquenal)."""
        db_dir = tmp_path / "db"
        repo = Repositorio("00000000000100", 2025, base_dir=db_dir)
        repo.criar_banco()

        destino = tmp_path / "parecer.docx"
        word_parecer.gerar(
            repo, 2025, tese="prescricao-quinquenal", destino=destino
        )

        assert destino.exists()
        doc = Document(str(destino))
        texto = _texto_concatenado(doc)
        assert "Prescrição Quinquenal" in texto
        assert "art. 168" in texto


class TestGerarSemAchados:
    def test_parecer_sem_oportunidades_gera_arquivamento(self, tmp_path):
        """Banco vazio + tese válida → doc gerado com seção de arquivamento."""
        db_dir = tmp_path / "db"
        repo = Repositorio("00000000000100", 2025, base_dir=db_dir)
        repo.criar_banco()

        destino = tmp_path / "parecer.docx"
        word_parecer.gerar(repo, 2025, tese="tema-69", destino=destino)

        assert destino.exists()
        doc = Document(str(destino))
        texto = _texto_concatenado(doc)

        # Identificação
        assert "PARECER TÉCNICO DE RECUPERAÇÃO TRIBUTÁRIA" in texto
        assert "00.000.000/0001-00" in texto
        assert "2025" in texto
        # Seção 5 com recomendação de arquivamento (sem achados)
        assert "arquivamento" in texto.lower()


class TestLogoOpcional:
    def test_sem_logo_usa_cabecalho_textual(self, tmp_path, monkeypatch):
        """Sem logo configurado, documento ainda é gerado com cabeçalho textual."""
        monkeypatch.setenv("PRIMETAX_LOGO_PATH", str(tmp_path / "inexistente.png"))
        db_dir = tmp_path / "db"
        repo = Repositorio("00000000000100", 2025, base_dir=db_dir)
        repo.criar_banco()

        destino = tmp_path / "parecer.docx"
        word_parecer.gerar(repo, 2025, tese="tema-69", destino=destino)
        assert destino.exists()

        doc = Document(str(destino))
        assert len(doc.inline_shapes) == 0  # nenhuma imagem embutida

    def test_com_logo_embed_picture(self, tmp_path, monkeypatch):
        """Logo PNG válido é embutido no documento (inline shape)."""
        from PIL import Image

        logo_path = tmp_path / "logo.png"
        Image.new("RGB", (40, 20), color=(0, 140, 149)).save(str(logo_path))
        monkeypatch.setenv("PRIMETAX_LOGO_PATH", str(logo_path))

        db_dir = tmp_path / "db"
        repo = Repositorio("00000000000100", 2025, base_dir=db_dir)
        repo.criar_banco()

        destino = tmp_path / "parecer.docx"
        word_parecer.gerar(repo, 2025, tese="tema-69", destino=destino)

        doc = Document(str(destino))
        assert len(doc.inline_shapes) == 1


class TestGerarComAchados:
    def test_parecer_tema69_com_cr07_gera_tabela(
        self, fixture_tese69_positivo, tmp_path
    ):
        """Fluxo real: importa EFD-Contrib com tese 69 positiva, roda diagnose,
        gera parecer tema-69 — deve conter CR-07 na tabela de resumo."""
        db_dir = tmp_path / "db"
        efd_contribuicoes.importar(
            fixture_tese69_positivo,
            encoding_override="utf8",
            prompt_operador=False,
            base_dir_db=db_dir,
        )
        repo = Repositorio("00000000000100", 2025, base_dir=db_dir)
        motor = Motor(repo)
        motor.diagnosticar_ano(2025)

        destino = tmp_path / "parecer.docx"
        word_parecer.gerar(repo, 2025, tese="tema-69", destino=destino)

        assert destino.exists()
        doc = Document(str(destino))
        texto = _texto_concatenado(doc)

        assert "Tese 69" in texto
        assert "CR-07" in texto
        assert "RE 574.706/PR" in texto
        assert "Análise Individualizada" in texto
        # Assinatura presente
        assert "Consultor Responsável" in texto

    def test_parecer_filtra_apenas_regras_da_tese(
        self, fixture_tese69_positivo, tmp_path
    ):
        """Parecer de 'insumos' não deve listar CR-07 (que é tema-69)."""
        db_dir = tmp_path / "db"
        efd_contribuicoes.importar(
            fixture_tese69_positivo,
            encoding_override="utf8",
            prompt_operador=False,
            base_dir_db=db_dir,
        )
        repo = Repositorio("00000000000100", 2025, base_dir=db_dir)
        motor = Motor(repo)
        motor.diagnosticar_ano(2025)

        destino = tmp_path / "parecer_insumos.docx"
        word_parecer.gerar(repo, 2025, tese="insumos", destino=destino)

        doc = Document(str(destino))
        texto = _texto_concatenado(doc)

        # Seção identifica a tese insumos
        assert "REsp 1.221.170" in texto
        # CR-07 não deve aparecer no resumo/análise desta tese
        assert "CR-07" not in texto
